from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
from m2x.Converter import get_config
from weasyprint import HTML
class CleanHTML():
    @classmethod
    def escape_html_code(self,html_content):
        import re
        content_unescaped = (
            html_content.replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("&quot;", '"')
            .replace("&#39;", "'")
        )
        code_pattern = re.compile(
            r'(<pre\s*[^>]*>\s*<code\s*[^>]*>)(.*?)(</code\s*>\s*</pre\s*>)',
            re.DOTALL | re.IGNORECASE
        )

        def replace_func(match):
            prefix = match.group(1)
            code = match.group(2)
            suffix = match.group(3)
            code_escaped = (
                code.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            return f"{prefix}{code_escaped}{suffix}"
        final_html = code_pattern.sub(replace_func, content_unescaped)
        extension = get_config.EXTENSION
        return extension + final_html
    

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

class HTMLToWordConverter:

    def _init_styles(self):
        normal = self.doc.styles['Normal']
        normal.font.name = 'Microsoft YaHei'
        normal.font.size = Pt(16)
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    def convert_html_to_word(self, html_content, save_path):
        self.doc = Document()
        self._init_styles()
        soup = BeautifulSoup(html_content, 'html.parser')
        for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre']):
            if tag.name in ['h1', 'h2', 'h3']:
                text = tag.get_text(strip=True)
                if text:
                    level = int(tag.name[1])
                    self.doc.add_heading(text, level=level)
            elif tag.name == 'p':
                text = tag.get_text(strip=True)
                if text:
                    self.doc.add_paragraph(text)

            elif tag.name in ['ul', 'ol']:
                for li in tag.find_all('li'):
                    t = li.get_text(strip=True)
                    if t:
                        style = 'List Bullet' if tag.name == 'ul' else 'List Number'
                        self.doc.add_paragraph(t, style=style)

            elif tag.name == 'pre':
                code = tag.get_text()
                if code.strip():
                    para = self.doc.add_paragraph()
                    run = para.add_run(code)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(45, 55, 72)
        self.doc.save(save_path)



class HTMLToPdfConverter():
    def convert_html_to_pdf(self,html_content,save_path):
            HTML(string=html_content).write_pdf(save_path)