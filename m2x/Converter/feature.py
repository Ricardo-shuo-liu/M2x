from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup

class HTMLToWordConverter:
    def __init__(self):
        self.doc = Document()
        self._init_styles()

    def _init_styles(self):
        """初始化Word样式"""
        # 普通文本样式
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Microsoft YaHei'
        normal_style.font.size = Pt(16)
        normal_style.font.color.rgb = RGBColor(51, 51, 51)
        normal_style.paragraph_format.line_spacing = 1.8
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 标题1样式
        h1_style = self.doc.styles['Heading 1']
        h1_style.font.name = 'Microsoft YaHei'
        h1_style.font.size = Pt(24)
        h1_style.font.color.rgb = RGBColor(44, 62, 80)
        self._add_paragraph_border(h1_style.paragraph_format, 'bottom', 'eeeeee')

        # 标题2样式
        h2_style = self.doc.styles['Heading 2']
        h2_style.font.name = 'Microsoft YaHei'
        h2_style.font.size = Pt(20)
        h2_style.font.color.rgb = RGBColor(44, 62, 80)
        self._add_paragraph_border(h2_style.paragraph_format, 'bottom', 'eeeeee')

    def _add_paragraph_border(self, para_format, border_pos, color_hex):
        """添加段落边框（底层XML操作）"""
        ppr = para_format._element
        p_border = OxmlElement('w:pBdr')
        border = OxmlElement(f'w:{border_pos}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:color'), color_hex)
        border.set(qn('w:sz'), '4')
        p_border.append(border)
        ppr.append(p_border)

    def _set_code_block_style(self, paragraph):
        """设置代码块样式"""
        # 背景色
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F8F9FA')
        paragraph._element.get_or_add_pPr().append(shading_elm)
        # 边框
        border_elm = OxmlElement('w:pBdr')
        for border_name in ['top', 'bottom', 'left', 'right']:
            b_elm = OxmlElement(f'w:{border_name}')
            b_elm.set(qn('w:val'), 'single')
            b_elm.set(qn('w:color'), 'E9ECEF')
            b_elm.set(qn('w:sz'), '4')
            border_elm.append(b_elm)
        paragraph._element.get_or_add_pPr().append(border_elm)
        # 间距
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_before = Pt(15)
        paragraph.paragraph_format.space_after = Pt(15)
        paragraph.paragraph_format.line_spacing = 1.5

    def _parse_list(self, list_tag):
        """解析列表（防None）"""
        if not list_tag:
            return
        list_type = list_tag.name
        for li in list_tag.find_all('li', recursive=False):
            if not li:
                continue
            para = self.doc.add_paragraph(style='List Bullet' if list_type == 'ul' else 'List Number')
            li_content = []
            for child in li.contents:
                if not child:
                    continue
                if hasattr(child, 'name') and child.name == 'p':
                    li_content.append(child.get_text(strip=True))
                elif hasattr(child, 'name') and child.name in ['ul', 'ol']:
                    self._parse_list(child)
                elif isinstance(child, str):
                    li_content.append(child.strip())
            if li_content:
                para.add_run(' '.join(li_content))

    def convert_html_to_word(self, html_content, save_path):
        """
        核心方法：HTML→Word（完善空值校验）
        :param html_content: 已转换的HTML字符串
        :param save_path: Word保存路径
        """
        # 1. 解析HTML（容错处理）
        soup = BeautifulSoup(html_content, 'html.parser')
        body = soup.body
        if not body:  # 防body为空
            raise ValueError("HTML内容中未找到<body>标签")

        # 2. 遍历节点（全量空值校验）
        for element in body.children:
            # 跳过None/空文本/注释节点
            if element is None or not hasattr(element, 'name'):
                continue
            tag_name = element.name
            if not tag_name:
                continue

            # 标题（h1-h6）
            if tag_name.startswith('h') and len(tag_name) == 2 and tag_name[1].isdigit():
                text = element.get_text(strip=True)
                if text:
                    self.doc.add_heading(text, level=int(tag_name[1]))

            # 普通段落
            elif tag_name == 'p':
                text = element.get_text(strip=True)
                if text:
                    self.doc.add_paragraph(text)

            # 列表
            elif tag_name in ['ul', 'ol']:
                self._parse_list(element)

            # 代码块
            elif tag_name == 'pre':
                code_text = element.get_text(strip=True)
                if code_text:
                    para = self.doc.add_paragraph()
                    self._set_code_block_style(para)
                    code_run = para.add_run(code_text)
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(14)
                    code_run.font.color.rgb = RGBColor(45, 55, 72)

        # 3. 保存文件
        self.doc.save(save_path)

# TODO:完成改部分逻辑阅读以及修复docx读取问题
